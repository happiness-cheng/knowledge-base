from app.utils.markdown_cleaner import clean_markdown, split_by_headings


def import_file(content: bytes, ext: str) -> list[tuple[str, str]]:
    if ext == "docx":
        return _import_docx(content)
    else:
        text = content.decode("utf-8", errors="replace")
        cleaned = clean_markdown(text)
        return split_by_headings(cleaned)


def _import_docx(content: bytes) -> list[tuple[str, str]]:
    from docx import Document
    import io

    doc = Document(io.BytesIO(content))

    sections = []
    current_title = "Untitled"
    current_lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if "Heading 2" in style:
            if current_lines:
                sections.append((current_title, "\n".join(current_lines)))
            current_title = text
            current_lines = []
        elif "Heading" in style:
            if current_lines:
                sections.append((current_title, "\n".join(current_lines)))
            current_title = text
            current_lines = []
        else:
            current_lines.append(text)

    if current_lines:
        sections.append((current_title, "\n".join(current_lines)))

    if not sections:
        all_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        all_text = "\n".join(all_paras)
        title = all_paras[0][:100] if all_paras else "Imported Document"
        sections = [(title, all_text)]

    return sections
